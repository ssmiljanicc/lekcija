#!/usr/bin/env python3
"""
export.py  –  Konverzija generisanih lekcija iz Markdown u Word (.docx) format.

═══════════════════════════════════════════════════════════════
 GĐDE SE POKREĆE
   Uvek iz korena projekta (folder gde se nalazi ovaj fajl):

       cd /Users/ssmiljanic/projekti/lekcija
       python3 export.py

 ULAZ — ODAKLE ČITA
   Čita .md fajlove iz foldera:

       output/           ← lekcije koje je generisao lekcija.py

   Ili prosleđuješ konkretan fajl:
       python3 export.py output/neka-lekcija.md

 IZLAZ — GDE UPISUJE
   .docx fajlovi se upisuju u:

       docs/             ← Word dokumenti, jedan po lekciji

 TIPIČAN REDOSLED KORAKA
   1.  python3 lekcija.py "input/shards/<fajl.md>"   # generiši lekciju → output/
   2.  python3 export.py                              # konvertuj → docs/
   3.  (opciono) python3 synthesize.py               # sintetiziraj sve lekcije

 OPCIJE
   --add-summary   Dodaj rezime u .md fajlove koji ga nemaju (bez API poziva)
   --md-only       Samo dodaj rezime, ne konvertuj u .docx
═══════════════════════════════════════════════════════════════
"""

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor


# --- Markdown parsing ---

def parse_markdown(text: str) -> list[dict]:
    """
    Parse markdown into a list of block elements.
    Each block: {"type": str, "content": str, "level": int}
    Types: heading, paragraph, code, bullet, numbered, blockquote, separator, frontmatter
    """
    # Strip YAML frontmatter
    text = re.sub(r"^---\n.*?\n---\n*", "", text, flags=re.DOTALL)

    blocks = []
    lines = text.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]

        # Fenced code block
        if line.startswith("```"):
            lang = line[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            blocks.append({"type": "code", "content": "\n".join(code_lines), "lang": lang})
            i += 1
            continue

        # Horizontal separator
        if re.match(r"^---+$", line.strip()):
            blocks.append({"type": "separator", "content": ""})
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.+)", line)
        if m:
            level = len(m.group(1))
            blocks.append({"type": "heading", "content": m.group(2).strip(), "level": level})
            i += 1
            continue

        # Blockquote
        if line.startswith("> "):
            content = line[2:]
            blocks.append({"type": "blockquote", "content": content})
            i += 1
            continue

        # Bullet list
        m = re.match(r"^(\s*)[*\-]\s+(.+)", line)
        if m:
            indent = len(m.group(1)) // 2
            blocks.append({"type": "bullet", "content": m.group(2).strip(), "level": indent})
            i += 1
            continue

        # Numbered list
        m = re.match(r"^(\s*)\d+\.\s+(.+)", line)
        if m:
            indent = len(m.group(1)) // 2
            blocks.append({"type": "numbered", "content": m.group(2).strip(), "level": indent})
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph - collect consecutive non-special lines
        para_lines = [line]
        i += 1
        while i < len(lines):
            next_line = lines[i]
            if (next_line.startswith("#") or next_line.startswith("```") or
                    next_line.startswith("> ") or re.match(r"^\s*[*\-]\s", next_line) or
                    re.match(r"^\s*\d+\.\s", next_line) or re.match(r"^---+$", next_line.strip()) or
                    not next_line.strip()):
                break
            para_lines.append(next_line)
            i += 1
        blocks.append({"type": "paragraph", "content": " ".join(para_lines)})

    return blocks


def apply_inline_formatting(run_adder, text: str):
    """
    Split text with **bold** and `code` markers and add runs with formatting.
    run_adder: callable(text, bold=False, code=False)
    """
    # Pattern: **bold**, `code`, or plain text
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    parts = pattern.split(text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run_adder(part[2:-2], bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run_adder(part[1:-1], code=True)
        elif part:
            run_adder(part)


# --- Word document builder ---

def set_heading_color(paragraph, hex_color: str):
    """Set paragraph font color via XML."""
    r_color = OxmlElement("w:color")
    r_color.set(qn("w:val"), hex_color)
    rPr = paragraph.runs[0]._r.get_or_add_rPr() if paragraph.runs else None
    if rPr is not None:
        rPr.append(r_color)


def add_horizontal_line(paragraph):
    """Add a horizontal line below a paragraph via XML border."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def build_docx(blocks: list[dict], output_path: Path, title: str = "", source: str = "", date_str: str = ""):
    doc = Document()

    # --- Page margins ---
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # --- Default body font ---
    style = doc.styles["Normal"]
    style.font.name = "Georgia"
    style.font.size = Pt(11)

    # --- Header with source info ---
    if source or date_str:
        header = doc.sections[0].header
        header_para = header.paragraphs[0]
        header_para.text = f"{source}  |  {date_str}" if (source and date_str) else (source or date_str)
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in header_para.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            run.font.name = "Arial"

    # --- Render blocks ---
    for block in blocks:
        btype = block["type"]

        if btype == "heading":
            level = block["level"]
            heading_style = f"Heading {min(level, 6)}"
            p = doc.add_heading("", level=level)
            p.clear()

            # Style mapping
            font_sizes = {1: 22, 2: 16, 3: 13, 4: 12, 5: 11, 6: 11}
            colors = {1: "1a1a2e", 2: "16213e", 3: "0f3460", 4: "333333", 5: "555555", 6: "777777"}

            def add_heading_run(text, bold=False, code=False):
                run = p.add_run(text)
                run.font.size = Pt(font_sizes.get(level, 11))
                run.font.color.rgb = RGBColor.from_string(colors.get(level, "000000"))
                run.font.name = "Arial" if level <= 2 else "Georgia"
                run.bold = bold or level <= 3
                if code:
                    run.font.name = "Courier New"
                    run.font.size = Pt(font_sizes.get(level, 11) - 1)

            apply_inline_formatting(add_heading_run, block["content"])

            if level == 2:
                add_horizontal_line(p)

            p.paragraph_format.space_before = Pt(18 if level == 1 else 14 if level == 2 else 10)
            p.paragraph_format.space_after = Pt(6)

        elif btype == "paragraph":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.space_before = Pt(2)

            def add_para_run(text, bold=False, code=False):
                run = p.add_run(text)
                run.bold = bold
                if code:
                    run.font.name = "Courier New"
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0xc7, 0x25, 0x4e)

            apply_inline_formatting(add_para_run, block["content"])

        elif btype == "code":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)

            # Gray background via shading XML
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "F4F4F4")
            pPr.append(shd)

            run = p.add_run(block["content"])
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x24, 0x29, 0x2e)

        elif btype == "bullet":
            level = block.get("level", 0)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
            p.paragraph_format.space_after = Pt(3)

            def add_bullet_run(text, bold=False, code=False):
                run = p.add_run(text)
                run.bold = bold
                if code:
                    run.font.name = "Courier New"
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0xc7, 0x25, 0x4e)

            apply_inline_formatting(add_bullet_run, block["content"])

        elif btype == "numbered":
            level = block.get("level", 0)
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
            p.paragraph_format.space_after = Pt(3)

            def add_num_run(text, bold=False, code=False):
                run = p.add_run(text)
                run.bold = bold
                if code:
                    run.font.name = "Courier New"
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0xc7, 0x25, 0x4e)

            apply_inline_formatting(add_num_run, block["content"])

        elif btype == "blockquote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)

            # Left border
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"), "single")
            left.set(qn("w:sz"), "12")
            left.set(qn("w:space"), "12")
            left.set(qn("w:color"), "4A90D9")
            pBdr.append(left)
            pPr.append(pBdr)

            run = p.add_run(block["content"])
            run.italic = True
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        elif btype == "separator":
            p = doc.add_paragraph()
            add_horizontal_line(p)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)

    doc.save(output_path)
    return output_path


# --- Summary injection for existing .md files without summary ---

def has_summary(text: str) -> bool:
    """Check if the document already has a summary/rezime section."""
    clean = re.sub(r"^---\n.*?\n---\n*", "", text, flags=re.DOTALL)
    lines = clean.strip().split("\n")

    # After H1, if the next non-empty line is another # heading or ---, no summary
    found_h1 = False
    for line in lines:
        if not line.strip():
            continue
        if line.startswith("# ") and not found_h1:
            found_h1 = True
            continue
        if found_h1:
            # If next content is a heading or separator - no summary
            if line.startswith("#") or line.startswith("---"):
                return False
            # Otherwise there's a paragraph after H1 = has summary
            return True
    return False


def extract_sections_summary(text: str) -> str:
    """Extract section headings and first sentence of each to build a local summary."""
    clean = re.sub(r"^---\n.*?\n---\n*", "", text, flags=re.DOTALL)
    sections = re.findall(r"^## (.+)", clean, re.MULTILINE)
    if not sections:
        return ""

    # Build a simple summary from section titles
    if len(sections) <= 3:
        topics = ", ".join(sections)
    else:
        topics = ", ".join(sections[:3]) + f" i još {len(sections)-3} tema"

    return (
        f"Ova lekcija pokriva sledeće teme: {topics}. "
        f"Dokument je organizovan u {len(sections)} sekcija koje zajedno pružaju "
        f"sveobuhvatan pregled tematike i praktičnih aspekata koji se obrađuju."
    )


def extract_key_conclusions(text: str) -> list[str]:
    """Extract bullet points from existing Ključni zaključci section, or return empty."""
    m = re.search(r"## Ključni zaključci\s*\n(.*?)(?=\n## |\Z)", text, re.DOTALL)
    if m:
        bullets = re.findall(r"^[-*]\s+(.+)", m.group(1), re.MULTILINE)
        return bullets
    return []


def inject_summary_into_md(md_path: Path) -> bool:
    """Add summary paragraph after H1 if missing. Returns True if modified."""
    text = md_path.read_text(encoding="utf-8")
    if has_summary(text):
        return False

    summary = extract_sections_summary(text)
    if not summary:
        return False

    # Find position after H1
    clean = re.sub(r"^---\n.*?\n---\n*", "", text, flags=re.DOTALL)
    h1_match = re.search(r"^# .+\n", clean, re.MULTILINE)
    if not h1_match:
        return False

    # Insert summary + separator after H1
    insert = f"\n{summary}\n\n---\n\n"
    new_clean = clean[:h1_match.end()] + insert + clean[h1_match.end():]

    # Re-attach frontmatter if present
    fm_match = re.match(r"^(---\n.*?\n---\n*)", text, re.DOTALL)
    if fm_match:
        new_text = fm_match.group(1) + new_clean
    else:
        new_text = new_clean

    md_path.write_text(new_text, encoding="utf-8")
    return True


# --- Main ---

def main():
    parser = argparse.ArgumentParser(
        description="Konvertuj markdown lekcije u Word (.docx) format."
    )
    parser.add_argument(
        "input_path",
        nargs="?",
        default="./output",
        help="Fajl ili folder sa .md fajlovima (default: ./output)",
    )
    parser.add_argument(
        "-o", "--output",
        default="./docs",
        help="Output folder za .docx fajlove (default: ./docs)",
    )
    parser.add_argument(
        "--add-summary",
        action="store_true",
        help="Dodaj rezime u .md fajlove koji ga nemaju (ne troši API pozive)",
    )
    parser.add_argument(
        "--md-only",
        action="store_true",
        help="Samo dodaj rezime u .md fajlove, bez konverzije u .docx",
    )

    args = parser.parse_args()

    input_path = Path(args.input_path)
    output_dir = Path(args.output)

    # Collect input files
    if input_path.is_dir():
        md_files = sorted(input_path.glob("*.md"))
    elif input_path.is_file() and input_path.suffix == ".md":
        md_files = [input_path]
    else:
        print(f"Error: {input_path} nije .md fajl ni folder.", file=sys.stderr)
        sys.exit(1)

    if not md_files:
        print(f"Nema .md fajlova u {input_path}")
        return

    # Step 1: Add summaries to .md files (no API calls)
    if args.add_summary or not args.md_only:
        print(f"Dodavanje rezimea u {len(md_files)} fajlova...")
        modified = 0
        for md in md_files:
            if inject_summary_into_md(md):
                print(f"  ✓ Dodat rezime: {md.name}")
                modified += 1
            else:
                print(f"  · Već ima rezime: {md.name}")
        print(f"  Izmenjeno: {modified}/{len(md_files)} fajlova\n")

    if args.md_only:
        return

    # Step 2: Convert to .docx
    output_dir.mkdir(parents=True, exist_ok=True)
    print(f"Konverzija u Word format → {output_dir}/")

    for md in md_files:
        text = md.read_text(encoding="utf-8")

        # Extract frontmatter metadata
        fm = re.match(r"^---\n(.*?)\n---", text, re.DOTALL)
        source = ""
        date_str = ""
        if fm:
            for line in fm.group(1).splitlines():
                if line.startswith("source:"):
                    source = line.split(":", 1)[1].strip()
                elif line.startswith("generated:"):
                    date_str = line.split(":", 1)[1].strip()

        blocks = parse_markdown(text)
        docx_path = output_dir / md.with_suffix(".docx").name

        build_docx(blocks, docx_path, source=source, date_str=date_str)
        print(f"  ✓ {md.name} → {docx_path.name}")

    print(f"\nGotovo. {len(md_files)} fajlova konvertovano u {output_dir}/")


if __name__ == "__main__":
    main()
